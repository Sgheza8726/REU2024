# split_paragraphs.py

import os
from docx import Document

def read_and_split(doc_path, output_dir):
    doc = Document(doc_path)
    os.makedirs(output_dir, exist_ok=True)
    total_paragraphs = len(doc.paragraphs)
    for i, paragraph in enumerate(doc.paragraphs):
        if paragraph.text.strip():
            paragraph_text = paragraph.text.strip()
            paragraph_file = os.path.join(output_dir, f'paragraph_{i+1}.txt')
            try:
                with open(paragraph_file, 'w', encoding='utf-8') as file:
                    file.write(paragraph_text)
            except Exception as e:
                print(f"Error writing paragraph {i+1} of {doc_path}: {e}")
        # Print progress every 100 paragraphs processed
        if (i + 1) % 100 == 0 or (i + 1) == total_paragraphs:
            print(f"Processed {i+1} / {total_paragraphs} paragraphs of {doc_path}")

def main():
    parts = ['Part_One.docx', 'Part_Two.docx', 'Part_Three.docx']
    output_dirs = ['Part_One_Paragraphs', 'Part_Two_Paragraphs', 'Part_Three_Paragraphs']

    for part, output_dir in zip(parts, output_dirs):
        print(f"Processing {part}...")
        read_and_split(part, output_dir)
        print(f"Finished processing {part}\n")

if __name__ == "__main__":
    main()

